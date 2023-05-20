import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.request import urlretrieve
import os

# Define the list of links to scrape
links = [
    'https://medicine.uiowa.edu/iowaprotocols/lower-lip-bump-mucocele-verus-mucus-retention-cyst-minor-salivary-gland-traumatic-lesion',
    'https://medicine.uiowa.edu/iowaprotocols/melanoma-evaluation-and-management-8th-edition-ajcc',
    'https://medicine.uiowa.edu/iowaprotocols/desmoplastic-melanoma',
    'https://medicine.uiowa.edu/iowaprotocols/melanoma-pathology-reporting-template',
    'https://medicine.uiowa.edu/iowaprotocols/squamous-cell-carcinoma-evaluation-and-management-historical-perspective',
    'https://medicine.uiowa.edu/iowaprotocols/laryngeal-leukoplakia-white-plaques-vocal-cords',
    'https://medicine.uiowa.edu/iowaprotocols/overview-squamous-dysplasia',
    'https://medicine.uiowa.edu/iowaprotocols/basosquamous-cell-cancer',
    'https://medicine.uiowa.edu/iowaprotocols/papillary-thyroid-carcinoma',
    'https://medicine.uiowa.edu/iowaprotocols/medullary-thyroid-carcinoma-pathology',
    'https://medicine.uiowa.edu/iowaprotocols/men-2b-multiple-endocrine-neoplasia-2b-sippels-syndrome',
    'https://medicine.uiowa.edu/iowaprotocols/thyroid-nodule-evaluation',
    'https://medicine.uiowa.edu/iowaprotocols/thyroidectomy-and-thyroid-lobectomy',
    'https://medicine.uiowa.edu/iowaprotocols/parathyroidectomy',
    'https://medicine.uiowa.edu/iowaprotocols/oropharyngeal-cancer-management',
    'https://medicine.uiowa.edu/iowaprotocols/oropharyngeal-cancer-hpv-and-patient-counseling',
    'https://medicine.uiowa.edu/iowaprotocols/transoral-robotic-surgery',
    'https://medicine.uiowa.edu/iowaprotocols/oral-cavity-and-oropharynx-protocols',
    'https://medicine.uiowa.edu/iowaprotocols/oral-cavity-and-oropharynx-protocols'
]
# Create the folder to store the documents
folder_name = "Iowa Head and Neck Protocols (uiowa.edu)/PDF"
os.makedirs(folder_name, exist_ok=True)

# Loop through each link and scrape its content
for link in links:
    # Make a GET request to the link
    response = requests.get(link)

    # Parse the HTML content using BeautifulSoup
    soup = BeautifulSoup(response.text, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Loop through each element in the HTML and add it to the document
    for element in soup.body.children:
        # If the element is a table, transform it into text with tab spaces
        if element.name == 'table':
            # Extract the table data and transform it into text with tab spaces
            table_text = ""
            for row in element.find_all('tr'):
                table_row = [cell.text for cell in row.find_all(['th', 'td'])]
                table_text += '\t'.join(table_row) + '\n'

            # Add the table text as a paragraph to the document
            doc.add_paragraph(table_text)

        # If the element is an image, download it and add it to the document
        elif element.name == 'img':
            # Get the source URL of the image
            src = element['src']

            # Download the image to a temporary file
            temp_file, _ = urlretrieve(src)

            # Add the image to the document
            doc.add_picture(temp_file)

            # Delete the temporary file
            os.remove(temp_file)

        # Otherwise, add the element's text to the document as a paragraph
        elif element.name is not None:
            doc.add_paragraph(element.text)

    # Save the document to disk within the folder
    filename = os.path.join(folder_name, link.split('/')[-1] + '.pdf')
    doc.save(filename)
